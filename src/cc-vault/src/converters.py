"""
Vault Converters - Convert various file formats to markdown for indexing

Supports:
- Word documents (.docx)
- PDF files (.pdf)
- Plain text (.txt)
- Markdown (.md) - passthrough
"""

import re
from pathlib import Path
from typing import Optional, Tuple

# Word document support (optional - required for .docx files)
DocxDocument = None
Table = None
Paragraph = None
try:
    from docx import Document as DocxDocument
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError:
    pass  # Will raise clear error when .docx conversion is attempted

# PDF support (optional - required for .pdf files)
fitz = None
try:
    import fitz  # pymupdf
except ImportError:
    pass  # Will raise clear error when .pdf conversion is attempted


# Supported file extensions
SUPPORTED_EXTENSIONS = {
    '.md': 'markdown',
    '.txt': 'text',
    '.docx': 'word',
    '.pdf': 'pdf',
}


def get_file_type(path: Path) -> Optional[str]:
    """Get the file type based on extension."""
    ext = path.suffix.lower()
    return SUPPORTED_EXTENSIONS.get(ext)


def is_supported(path: Path) -> bool:
    """Check if a file type is supported for conversion."""
    return get_file_type(path) is not None


def convert_to_markdown(path: Path) -> Tuple[str, dict]:
    """
    Convert a file to markdown format.

    Args:
        path: Path to the file to convert

    Returns:
        Tuple of (markdown_content, metadata)
        metadata includes: original_format, page_count (for PDFs), etc.

    Raises:
        ValueError: If file type is not supported
        FileNotFoundError: If file doesn't exist
    """
    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    file_type = get_file_type(path)
    if file_type is None:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    if file_type == 'markdown':
        return _convert_markdown(path)
    elif file_type == 'text':
        return _convert_text(path)
    elif file_type == 'word':
        return _convert_word(path)
    elif file_type == 'pdf':
        return _convert_pdf(path)
    else:
        raise ValueError(f"No converter for type: {file_type}")


def _convert_markdown(path: Path) -> Tuple[str, dict]:
    """Passthrough for markdown files."""
    content = path.read_text(encoding='utf-8')
    return content, {'original_format': 'markdown'}


def _convert_text(path: Path) -> Tuple[str, dict]:
    """Convert plain text to markdown."""
    content = path.read_text(encoding='utf-8')
    # Wrap in code block if it looks like code, otherwise return as-is
    return content, {'original_format': 'text'}


def _convert_word(path: Path) -> Tuple[str, dict]:
    """
    Convert Word document (.docx) to markdown.

    Preserves:
    - Headings (converted to # syntax)
    - Bold and italic text
    - Bullet and numbered lists
    - Tables (converted to markdown tables)
    - Basic paragraph structure
    """
    if DocxDocument is None:
        raise ImportError(
            "python-docx is required for Word document conversion.\n"
            "Install with: pip install python-docx"
        )

    doc = DocxDocument(str(path))
    markdown_lines = []
    metadata = {
        'original_format': 'word',
        'paragraph_count': len(doc.paragraphs),
        'table_count': len(doc.tables),
    }

    # Process document elements in order
    for element in doc.element.body:
        if element.tag.endswith('p'):
            # Paragraph
            para = Paragraph(element, doc)
            md_line = _convert_paragraph(para)
            if md_line:
                markdown_lines.append(md_line)
        elif element.tag.endswith('tbl'):
            # Table
            table = Table(element, doc)
            md_table = _convert_table(table)
            if md_table:
                markdown_lines.append(md_table)

    content = '\n\n'.join(markdown_lines)
    return content, metadata


def _convert_paragraph(para: 'Paragraph') -> str:
    """Convert a Word paragraph to markdown."""
    text = para.text.strip()
    if not text:
        return ''

    style_name = para.style.name.lower() if para.style else ''

    # Handle headings
    if 'heading 1' in style_name:
        return f'# {text}'
    elif 'heading 2' in style_name:
        return f'## {text}'
    elif 'heading 3' in style_name:
        return f'### {text}'
    elif 'heading 4' in style_name:
        return f'#### {text}'
    elif 'heading' in style_name:
        return f'##### {text}'
    elif 'title' in style_name:
        return f'# {text}'
    elif 'subtitle' in style_name:
        return f'*{text}*'

    # Handle list items
    if 'list' in style_name or 'bullet' in style_name:
        return f'- {text}'
    elif para.style.name.startswith('List Number'):
        return f'1. {text}'

    # Process inline formatting
    formatted_text = _format_runs(para)

    return formatted_text


def _format_runs(para: 'Paragraph') -> str:
    """Process runs within a paragraph for bold/italic formatting."""
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue

        if run.bold and run.italic:
            text = f'***{text}***'
        elif run.bold:
            text = f'**{text}**'
        elif run.italic:
            text = f'*{text}*'

        parts.append(text)

    return ''.join(parts)


def _convert_table(table: 'Table') -> str:
    """Convert a Word table to markdown table."""
    if not table.rows:
        return ''

    md_lines = []

    for i, row in enumerate(table.rows):
        cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
        md_lines.append('| ' + ' | '.join(cells) + ' |')

        # Add header separator after first row
        if i == 0:
            separator = '| ' + ' | '.join(['---'] * len(cells)) + ' |'
            md_lines.append(separator)

    return '\n'.join(md_lines)


def _convert_pdf(path: Path) -> Tuple[str, dict]:
    """
    Convert PDF to markdown.

    Uses pymupdf for text extraction. Handles:
    - Multi-page documents
    - Basic text structure
    - Page breaks as horizontal rules
    """
    if fitz is None:
        raise ImportError(
            "pymupdf is required for PDF conversion.\n"
            "Install with: pip install pymupdf"
        )

    doc = fitz.open(str(path))
    markdown_pages = []
    metadata = {
        'original_format': 'pdf',
        'page_count': len(doc),
        'title': doc.metadata.get('title', ''),
        'author': doc.metadata.get('author', ''),
    }

    for page_num, page in enumerate(doc):
        # Extract text with basic structure preservation
        text = page.get_text("text")

        if text.strip():
            # Clean up the text
            text = _clean_pdf_text(text)

            # Add page marker for multi-page docs
            if len(doc) > 1:
                markdown_pages.append(f'<!-- Page {page_num + 1} -->\n\n{text}')
            else:
                markdown_pages.append(text)

    doc.close()

    # Join pages with horizontal rules
    content = '\n\n---\n\n'.join(markdown_pages)

    return content, metadata


def _clean_pdf_text(text: str) -> str:
    """Clean up extracted PDF text."""
    lines = text.split('\n')
    cleaned_lines = []
    prev_line_empty = False

    for line in lines:
        line = line.rstrip()

        # Skip multiple consecutive empty lines
        if not line:
            if not prev_line_empty:
                cleaned_lines.append('')
                prev_line_empty = True
            continue

        prev_line_empty = False

        # Detect potential headings (ALL CAPS lines, short lines)
        if line.isupper() and len(line) < 100:
            line = f'## {line.title()}'

        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines)


def convert_and_save(
    source_path: Path,
    output_dir: Optional[Path] = None,
    output_name: Optional[str] = None
) -> Tuple[Path, str, dict]:
    """
    Convert a file and save as markdown.

    Args:
        source_path: Path to source file
        output_dir: Directory for output (default: same as source)
        output_name: Output filename without extension (default: same as source)

    Returns:
        Tuple of (output_path, content, metadata)
    """
    source_path = Path(source_path)

    if output_dir is None:
        output_dir = source_path.parent
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    if output_name is None:
        output_name = source_path.stem

    output_path = output_dir / f'{output_name}.md'

    # Convert
    content, metadata = convert_to_markdown(source_path)

    # Save
    output_path.write_text(content, encoding='utf-8')

    return output_path, content, metadata
