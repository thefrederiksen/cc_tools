"""Word document conversion using python-docx."""

from pathlib import Path
from typing import Optional

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def convert_to_word(html_content: str, output_path: Path) -> None:
    """
    Convert HTML to Word document.

    Maps HTML elements to Word styles:
    - h1-h6 -> Heading 1-6
    - p -> Normal
    - ul/ol -> List styles
    - table -> Table
    - pre/code -> Code style
    - blockquote -> Quote style

    Args:
        html_content: Complete HTML document string
        output_path: Path for output .docx file
    """
    # Parse HTML
    soup = BeautifulSoup(html_content, "html.parser")

    # Create document
    doc = Document()

    # Find the main content
    body = soup.find("article") or soup.find("body") or soup

    # Process elements
    _process_element(doc, body)

    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)

    # Save
    doc.save(str(output_path))


def _process_element(doc: Document, element, parent_list: Optional[str] = None):
    """Recursively process HTML elements."""
    if element.name is None:
        # Text node
        return

    for child in element.children:
        if child.name is None:
            # Skip whitespace-only text
            if child.string and child.string.strip():
                # Text should be handled by parent element
                pass
            continue

        if child.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
            level = int(child.name[1])
            text = child.get_text(strip=True)
            doc.add_heading(text, level=level)

        elif child.name == "p":
            text = child.get_text(strip=True)
            if text:
                doc.add_paragraph(text)

        elif child.name == "ul":
            _process_list(doc, child, ordered=False)

        elif child.name == "ol":
            _process_list(doc, child, ordered=True)

        elif child.name == "blockquote":
            text = child.get_text(strip=True)
            if text:
                para = doc.add_paragraph(text)
                para.style = "Quote" if "Quote" in [s.name for s in doc.styles] else "Normal"
                para.paragraph_format.left_indent = Inches(0.5)

        elif child.name == "pre":
            code_text = child.get_text()
            para = doc.add_paragraph()
            run = para.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            para.paragraph_format.left_indent = Inches(0.25)

        elif child.name == "table":
            _process_table(doc, child)

        elif child.name == "hr":
            # Add horizontal line as paragraph with border
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(12)

        elif child.name in ["div", "article", "section", "main"]:
            # Container elements - recurse
            _process_element(doc, child)


def _process_list(doc: Document, list_element, ordered: bool = False):
    """Process ul or ol list."""
    for li in list_element.find_all("li", recursive=False):
        text = li.get_text(strip=True)
        if text:
            style = "List Number" if ordered else "List Bullet"
            # Fallback if style doesn't exist
            try:
                doc.add_paragraph(text, style=style)
            except KeyError:
                para = doc.add_paragraph(text)
                if ordered:
                    para.paragraph_format.left_indent = Inches(0.5)
                else:
                    para.paragraph_format.left_indent = Inches(0.5)


def _process_table(doc: Document, table_element):
    """Process HTML table."""
    rows = table_element.find_all("tr")
    if not rows:
        return

    # Count columns from first row
    first_row = rows[0]
    cols = first_row.find_all(["th", "td"])
    num_cols = len(cols)

    if num_cols == 0:
        return

    # Create table
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    # Fill table
    for row_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for col_idx, cell in enumerate(cells):
            if col_idx < num_cols:
                text = cell.get_text(strip=True)
                table.rows[row_idx].cells[col_idx].text = text

                # Bold headers
                if cell.name == "th":
                    for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
